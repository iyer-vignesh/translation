from django.shortcuts import render
from django.http import HttpResponse
from .models import folder
# from . import final_translation

# Create your views here.

import docx
import PyPDF2
from translate import Translator
import requests
from docx2pdf import convert
from docx.enum.text import WD_ALIGN_PARAGRAPH



def readtxt(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)


def readpdf(pdfName):
  read_pdf = PyPDF2.PdfReader(pdfName)
  fulltext = ''

  for i in range(0, len(read_pdf.pages)):
      page = read_pdf.pages[i]
      page_content = page.extract_text()
      fulltext += page_content
  return fulltext

translator= Translator(from_lang= 'en', to_lang="ar")


def trans(text):
    marker_1 = True
    marker_2 = True

    doc = text.split('\n')  # Split text paragraph wise & have them as list
    para = [i.split('.') for i in doc]  # Sentence tokenize each paragraph & keep them in separate lists

    corpus = []

    for sentences in para:

        multi_sent = []
        single_sent = []

        if len(sentences) > 1:

            for i in sentences:

                if len(i) > 499:
                    ind = len(i) // 2
                    while marker_1 == True:
                        ind += 1
                        if i[ind] == ' ':
                            part_1 = i[:ind + 1]
                            part_2 = i[ind:]
                            marker_1 = False

                    sent_1 = translator.translate(part_1)
                    sent_2 = translator.translate(part_2)
                    sent_3 = sent_1 + ' ' + sent_2
                    multi_sent.append(sent_3)


                else:
                    txt = translator.translate(i)
                    multi_sent.append(txt)

            corpus.append(multi_sent)

        elif len(sentences) == 1:

            sent_0 = sentences[0]

            if len(sent_0) > 499:
                ind = len(sent_0) // 2
                while marker_2 == True:
                    ind += 1
                    if i[ind] == ' ':
                        part_0_1 = sent_0[:ind + 1]
                        part_0_2 = sent_0[ind:]
                        marker_2 = False

                sent_0_1 = translator.translate(part_0_1)
                sent_0_2 = translator.translate(part_0_2)
                sent_0_3 = sent_0_1 + ' ' + sent_0_2
                single_sent.append(sent_0_3)


            else:
                txt_0 = translator.translate(sent_0)
                single_sent.append(txt_0)

            corpus.append(single_sent)


        else:
            continue

    return corpus

def write_file(text):
  global trans_file
  trans_text = []
  translated_file = 'Translated_file.docx'
  for para in text:
    temp = ''

    if len(para) > 1:
      for sent in para:
        temp += sent
      trans_text.append([temp])

    else:
      trans_text.append(para)

  # Create a new docx file
  document = docx.Document()
  # Add a heading
  # document.add_heading(trans_text[0][0], level=0)

  heading = document.add_heading(trans_text[0][0], level=0)
  heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
  paragraph = document.add_paragraph()
  paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

  for para in trans_text[1:]:
    # paragraph = document.add_paragraph(para)

    paragraph.add_run(para)

  document.save(translated_file)
  # Converting docx present in the same folder
  # as the python file
  trans_file = convert(translated_file)

  return trans_file

def create_file(file_name):

    # link = 'https://translationhosting.pythonanywhere.com/media/Desktop/'
    link = "http://127.0.0.1:8000/media/Desktop/"
    file_format = file_name.split('.')[-1]

    if file_format == 'docx':
        new_file = "test.docx"

        response = requests.get(link+file_name)
        open(new_file, "wb").write(response.content)

    elif file_format == 'pdf':
        new_file = "test.pdf"

        response = requests.get(link+file_name)
        open(new_file, "wb").write(response.content)

    else:
        print('Unsupported File Type')

    return new_file

def final_translator(path):

  path_list = path.split('.')
  file_type = path_list[-1]
  txt = ''
  marker = False

  if file_type == 'docx':
    txt = readtxt(path)

  elif file_type == 'pdf':
    txt = readpdf(path)


  else:
    marker = True

  if not marker:

    # lang = detect(txt)
    print('The uploaded file type is ', file_type)

    translation = trans(txt)

  else:
    print('Unsupported File Type')

  final_trans_file = write_file(translation)

  return final_trans_file




def demo(request):
    global name
    if request.method == 'POST':
        upload=request.FILES['abc']
        fol=folder(upload=upload)
        fol.save()
    name=folder.objects.all().last()
    return render(request, 'new.html',{'name':name})


# def translate_test(request):
#     global final_file
#     global file_name
#
#     file = str(name)
#     file_name = file.split('/')[-1]
#     # print(fii.split('/')[-1])
#     create_file(file_name)
#
#     file_format = file_name.split('.')[-1]
#
#     if file_format == 'docx':
#         f_name = "test.docx"
#
#     elif file_format == 'pdf':
#         f_name = "test.pdf"
#
#     path = f_name
#     final_file = final_translator(path)
#     response = HttpResponse(content_type='application/pdf')
#     response['Content-Disposition'] = "attachment; filename= file.pdf "
#     print(type(trans_file))
#     return response

from django.http import FileResponse


def translate_test(request):

    global final_file
    global file_name

    file = str(name)
    file_name = file.split('/')[-1]
    # print(fii.split('/')[-1])
    create_file(file_name)

    file_format = file_name.split('.')[-1]

    if file_format == 'docx':
        f_name = "test.docx"

    elif file_format == 'pdf':
        f_name = "test.pdf"

    path = f_name
    final_file = final_translator(path)

    file = open('C:/Users/venka/Downloads/ERP-PRODUCT/ERP-PRODUCT/pythonProject/demoproject/Translated_file.pdf', 'rb')
    response = FileResponse(file)
    response['Content-Type'] = 'application/pdf'
    response['Content-Disposition'] = 'attachment; filename="file.pdf"'
    return response


import mimetypes

# def download_file(request):
#     # fill these variables with real values
#     # fl_path = final_file
#     # filename = 'Translated '+ file_name +'.pdf'
#
#     # fl = open(fl_path, 'r')
#     # mime_type, _ = mimetypes.guess_type(fl_path)
#     response = HttpResponse(content_type= 'application/pdf')
#     # response['Content-Disposition'] = "attachment; filename=%s" % filename
#     response['Content-Disposition'] = "attachment; filename= '%s'" % final_file
#     print(type(response))
#
#     return response
"D:\Oracuz\erp\ERP-PRODUCT\pythonProject\demoproject\Translated_file.pdf"